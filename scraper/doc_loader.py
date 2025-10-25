from typing import Dict, List
import docx
import os
import json

def read_word_document(file_path: str) -> str:
    """
    Read content from a Word document.
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        str: The text content of the document
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found at {file_path}")
        
    if not file_path.endswith('.docx'):
        raise ValueError("File must be a .docx document")
    
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n'.join(full_text)

def ingest_word_document(file_path: str, cache_file: str = 'scraped_data.json') -> None:
    """
    Ingest content from a Word document and add it to the cache.
    
    Args:
        file_path (str): Path to the Word document
        cache_file (str): Path to the cache file
    """
    try:
        # Read the document
        content = read_word_document(file_path)
        
        # Create a document entry
        doc_entry = {
            'url': f"file://{os.path.abspath(file_path)}",
            'content': content,
            'source': 'word_document',
            'title': os.path.basename(file_path)
        }
        
        # Load existing cache
        if os.path.exists(cache_file):
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
        else:
            cache_data = []
            
        # Add or update document entry
        updated = False
        for i, entry in enumerate(cache_data):
            if entry.get('url') == doc_entry['url']:
                cache_data[i] = doc_entry
                updated = True
                break
                
        if not updated:
            cache_data.append(doc_entry)
            
        # Save updated cache
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=2)
            
    except Exception as e:
        raise Exception(f"Error ingesting Word document: {str(e)}")